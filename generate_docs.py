import base64
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    return p

def add_mermaid_diagram(doc, mermaid_code, caption):
    # Encode for mermaid.ink
    try:
        # Base64 urlsafe encode the mermaid code
        encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
        url = f"https://mermaid.ink/img/{encoded}?type=png"
        
        response = requests.get(url)
        if response.status_code == 200:
            image_stream = BytesIO(response.content)
            doc.add_picture(image_stream, width=Inches(6.0))
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph(f"[Failed to load diagram from API. Status: {response.status_code}]")
    except Exception as e:
        doc.add_paragraph(f"[Error generating diagram: {str(e)}]")

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Spendora Advanced Financial AI System - Project Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ----------- SECTION 1: INTRODUCTION -----------
    add_heading(doc, '1. INTRODUCTION', level=1)
    
    add_heading(doc, '1.1 Existing System', level=2)
    add_bullet(doc, "The existing manual financial tracking systems predominantly rely on physical ledger books, disjointed spreadsheet software, or outdated monolithic desktop applications that lack mobility, interconnectivity, and real-time synchronization. These legacy systems require intense manual data entry, creating a high barrier to entry and a significant probability of human error when transcribing receipts or logging daily expenditures.")
    add_bullet(doc, "The vast majority of current solutions do not integrate modern artificial intelligence to minimize the administrative burden of classifying transactions. Consequently, users are often discouraged from maintaining long-term financial discipline due to the sheer volume of manual input required, which frequently results in lost records, mismatched budgeting calculations, and overall poor visibility into personal aggregate financial health.")
    add_bullet(doc, "Furthermore, without predictive modeling, most systems serve merely as reactive historical repositories rather than proactive platforms that can anticipate upcoming bills, alert the user about budget exhaustion before it happens, or provide dynamic insights into spending velocity throughout a given month.")

    add_heading(doc, '1.2 Need for System', level=2)
    add_bullet(doc, "The modern individual faces an unprecedented volume of micro-transactions, utility bills, digital subscriptions, and varied income streams, creating an absolute necessity for an omnipotent, automated financial concierge. The need for Spendora arises specifically from the deficit of tools that understand context and reduce friction in daily accounting, turning an arduous chore into an interactive, instantaneous process via voice and visual recognition.")
    add_bullet(doc, "Users require an intelligent assistant capable of extracting unstructured data from physical receipts instantly to securely document expenses, circumventing the manual logbook process entirely. Similarly, vocal intake speeds up on-the-go tracking, directly addressing the core problem of user retention in daily financial software.")
    add_bullet(doc, "By enforcing hard limits via interactive budgeting modules and sending proactive notifications prior to exceeding expenditure thresholds, individuals can actively safeguard their liquid assets. The system serves the critical need of converting passive financial observation into active economic defense, allowing households to aggressively identify wasteful spending, systematically manage recurring obligations, and securely track the growth of their specialized savings pools.")

    add_heading(doc, '1.3 Scope of Project', level=2)
    add_bullet(doc, "The operational boundary of the Spendora project encompasses a massive suite of features strictly aimed at personal, localized, and multi-faceted financial tracking. The core scope involves the establishment of a robust relational data model managing User Profiles, localized Categories, independent ledgers for Income and Expenses, specialized Savings Vaults, and an automated Recurring Bill schedule engine.")
    add_bullet(doc, "Its functional scope strictly integrates the external Google Gemini Artificial Intelligence endpoints strictly for the natural language processing of audio commands and the Optical Character Recognition processing of receipt imaging. The application does not natively execute banking transactions or interface with external financial institution APIs (e.g., Plaid); rather, it acts as the centralized personal record and analytics layer.")
    add_bullet(doc, "The architectural scope mandates the delivery of a responsive, pristine, web-accessible dashboard rendering advanced HTML5 and vanilla JavaScript analytics charts for immediate trend visualization. It handles all authentication locally using rigorous cryptographic hashing and includes internal account recovery modules independent of third-party mailing protocols.")

    add_heading(doc, '1.4 Operating Environment', level=2)
    add_bullet(doc, "Hardware Requirements:")
    add_bullet(doc, "\t- Standard consumer-grade web-capable device (Desktop Computer, Laptop, Tablet, or Smartphone) equipped with at least 2GB of RAM for modern browser execution.")
    add_bullet(doc, "\t- Functioning audio microphone device for Voice Command operational input layers and an optical camera (or image upload capability) for Receipt Scanning features.")
    add_bullet(doc, "\t- Server-side hardware configuration requiring at minimum a dual-core CPU processing module, 2GB internal RAM, and 10GB non-volatile storage to comfortably execute the Python environment runtime and host the embedded localized SQL database efficiently.")
    add_bullet(doc, "Software Requirements:")
    add_bullet(doc, "\t- Frontend Technologies: Vanilla HTML5 for foundational semantic tree structuring, Custom CSS3 for utilizing dynamic flexbox/grid alignments natively supporting animations/glassmorphism, and standard ES6 Vanilla JavaScript strictly to handle asynchronous fetch commands and manipulate client-side Chart.js visual renders natively.")
    add_bullet(doc, "\t- Backend Technology: The Python 3.10+ programming language environment executing the highly robust, scalable, Model-View-Template synchronized framework ecosystem of Django 4.")
    add_bullet(doc, "\t- Database: Relational schema management currently implemented through SQLite3 for seamless localized deployment and state-machine transitions, structurally capable of immediate migration execution to an enterprise PostgreSQL server instance.")
    add_bullet(doc, "\t- Web Server: The native internal WSGI standard development daemon provided inherently by the Django platform, fully prepared for deployment transitioning onto a production-grade external Apache or Nginx layer stack via Gunicorn binding protocols.")


    # ----------- SECTION 2: PROPOSED SYSTEM -----------
    add_heading(doc, '2. PROPOSED SYSTEM', level=1)
    
    add_heading(doc, '2.1 Objectives of the System', level=2)
    add_bullet(doc, "Fundamentally eliminate friction during the daily fiscal entry sequence by providing zero-touch tracking methods natively supported across mobile and desktop interfaces, leveraging deep learning vision OCR models and advanced voice classification engines to categorize and calculate expenditures in real-time.")
    add_bullet(doc, "Synthesize complete financial clarity by creating an overarching intuitive dashboard system that visualizes complicated cash-flow movements intuitively. It aggregates diverse data arrays—Incomes, Expenses, Savings, and Bill Outflows—and plots this information across interactive visual elements, ultimately exposing otherwise hidden monthly consumption trends.")
    add_bullet(doc, "Mechanize proactive defense against chronic overspending through tight integrations of an advanced budgeting limitation algorithm. By measuring category-specific expenses against localized constraints, our objective is to intelligently trigger frontend UI blockades and dashboard alerts the moment eighty percent budget threshold consumption is identified.")
    
    add_heading(doc, '2.2 User Requirement', level=2)
    add_bullet(doc, "The end-user inherently requires a protected platform implementing sophisticated password hashing encryption and an accessible independent recovery mechanism through a localized, private Security Question/Answer vector avoiding complex dependency architectures.")
    add_bullet(doc, "They require an ecosystem capable of retaining customized categorizations so they can separate non-tangible expenditures effectively, coupled directly with the ability to define distinct, temporal spending ceilings (Budgets) independently across each personalized category.")
    add_bullet(doc, "There is a strict user requirement for an automated billing component that maintains memory of cyclical financial obligations—such as monthly mortgage, utility packages, and internet subscriptions—triggering visible pre-due dashboard notifications and optionally recreating the bill for the subsequent month cycle automatically upon user markup.")

    # ----------- SECTION 3: ANALYSIS AND DESIGN -----------
    add_heading(doc, '3. ANALYSIS AND DESIGN', level=1)

    add_heading(doc, '3.1. ER Diagram', level=2)
    add_bullet(doc, "The Entity-Relationship visual model illustrates the distinct normalized constraints outlining our relational storage architecture securely. It connects the overarching User entity explicitly via one-to-many cardinality vectors cascading outwards towards distinct modular financial logs, Categories, and Budget constraints.")
    
    erd_mermaid = "erDiagram\n    USER ||--o{ USER_PROFILE : has\n    USER ||--o{ CATEGORY : creates\n    USER ||--o{ EXPENSE : logs\n    USER ||--o{ INCOME : earns\n    USER ||--o{ SAVING : saves\n    USER ||--o{ BUDGET : sets\n    USER ||--o{ BILL_REMINDER : schedules\n    CATEGORY ||--o{ EXPENSE : categorizes\n    CATEGORY ||--o{ BUDGET : limits"
    add_mermaid_diagram(doc, erd_mermaid, "Figure 3.1: ER Diagram")

    add_heading(doc, '3.1.1 DFD', level=2)
    add_bullet(doc, "This Data Flow Level visualization explicitly details how primary environmental payloads—such as graphical interface submissions and raw audio vocal feeds—are processed sequentially through the Application controller, routing securely through the external asynchronous AI interpretation module, before safely integrating downstream into finalized structured Database storage records.")

    dfd_mermaid = "flowchart TD\n    USR((User))\n    sys{Spendora \nApplication}\n    AI((AI API))\n    DB[(Database)]\n\n    USR -- Audio/Images/Data --> sys\n    sys -- Dashboard Output --> USR\n    sys -- API Payload --> AI\n    AI -- Parsed JSON --> sys\n    sys -- SQL Insert/Query --> DB\n    DB -- Record Matrix --> sys"
    add_mermaid_diagram(doc, dfd_mermaid, "Figure 3.1.1: Context DFD (Level 0)")

    add_heading(doc, '3.2. Use Case Diagram', level=2)
    add_bullet(doc, "The Use Case schematic abstracts internal complexities, depicting purely the primary actionable capabilities an actor naturally expects when engaging the Spendora system environment. It demonstrates operational autonomy mapping registration, voice transactions, bill management, and visual analytics interaction explicitly to client-side initiation parameters.")

    use_case_mermaid = "flowchart LR\n    User((Actor))\n    UC1([Register System])\n    UC2([Voice Submits])\n    UC3([Scan Receipt AI])\n    UC4([Manage Budget Limit])\n    UC5([Automate Bills])\n    \n    User --> UC1\n    User --> UC2\n    User --> UC3\n    User --> UC4\n    User --> UC5"
    add_mermaid_diagram(doc, use_case_mermaid, "Figure 3.2: Use Case Diagram")

    add_heading(doc, '3.3. Activity Diagram', level=2)
    add_bullet(doc, "This Activity algorithmic flow traces the exact mathematical and state-machine logic required sequentially during an expense addition runtime. We expose decision divergence regarding data validity and aggressively trigger independent state shifts if the current transactional sum cascades over predefined eighty-percent budget threshold alarms.")

    act_mermaid = "stateDiagram-v2\n    [*] --> FormInput\n    FormInput --> Validate\n    Validate --> DB_Insert : True\n    Validate --> FormInput : False\n    DB_Insert --> QueryBudget\n    QueryBudget --> CheckLimit\n    CheckLimit --> ProcessNormal: <80%\n    CheckLimit --> TriggerWarning: >=80%\n    TriggerWarning --> [*]\n    ProcessNormal --> [*]"
    add_mermaid_diagram(doc, act_mermaid, "Figure 3.3: Activity Diagram")

    add_heading(doc, '3.4. Sequence Diagram', level=2)
    add_bullet(doc, "The sequence model explicitly describes our most intricate synchronization path: The Receipt AI Scanning execution flow. It details chronological object messages traversing the HTML Frontend, resolving via Django WSGI handlers, transmitting Base64 graphics over external TLS API tunnels, and parsing JSON callbacks structurally into relational category assignments.")

    seq_mermaid = "sequenceDiagram\n    actor U as User\n    participant V as Interface\n    participant C as Controller\n    participant G as Gemini AI\n    participant D as Database\n\n    U->>V: Upload Image & Request Scan\n    V->>C: POST Multipart File\n    C->>G: Request OCR JSON Schema\n    G-->>C: Response Payload\n    C->>D: Verify Category OR Insert\n    D-->>C: Object ID\n    C->>D: Save Expense Data\n    C-->>V: Return Success HTTP\n    V-->>U: Render Status"
    add_mermaid_diagram(doc, seq_mermaid, "Figure 3.4: Sequence Diagram")

    add_heading(doc, '3.5. Class Diagram', level=2)
    add_bullet(doc, "The analytical Object Class visualization abstracts the Spendora Python Model code base directly. It identifies inheritance hierarchies, mapping common functionality underneath an abstract transaction interface, while maintaining modular independent class variables mapping our internal logical controllers natively representing Django's standardized models implementation.")

    class_mermaid = "classDiagram\n    class User {\n        +id: int\n        +username: string\n        +login()\n    }\n    class Transaction {\n        <<Abstract>>\n        +amount: decimal\n        +date: date\n        +save()\n    }\n    class Expense {\n        +description: string\n    }\n    class Saving {\n        +source: string\n    }\n    User \"1\" --> \"*\" Transaction\n    Transaction <|-- Expense\n    Transaction <|-- Saving"
    add_mermaid_diagram(doc, class_mermaid, "Figure 3.5: Class Diagram")

    add_heading(doc, '3.6 Package Diagram', level=2)
    add_bullet(doc, "The package breakdown architecture organizes the Django foundational layer namespaces. It logically decouples administrative system configurations from dynamic application functionality, guaranteeing that our specific 'Tracker' application handles all temporal financial routing natively separated from the static rendering assets and framework utilities.")

    pkg_mermaid = "flowchart TD\n    subgraph expense_ai_system [Root Project]\n        A[manage.py Core]\n        subgraph Config Layer\n            B[settings.py]\n            C[urls.py]\n        end\n        subgraph Tracker App\n            D[models.py]\n            E[views.py]\n            F[forms.py]\n        end\n        subgraph Interface\n            G[static/ CSS-JS]\n            H[templates/ HTML]\n        end\n    end\n    Config Layer --> Tracker App\n    Tracker App --> Interface"
    add_mermaid_diagram(doc, pkg_mermaid, "Figure 3.6: Package Diagram")


    add_heading(doc, '3.7 Component Diagram', level=2)
    add_bullet(doc, "This macro Component view exposes macro-structural subsystems and their dependency interaction protocols. The HTTP/WS client interactions bridge towards our secure Controller backend, which independently delegates heavy asynchronous workloads seamlessly out towards cloud intelligent processing environments while simultaneously interacting deeply with standard localized object persistence stores.")

    comp_mermaid = "flowchart LR\n    UI[Frontend UI] <-->|Requests| API[Django Views]\n    API <-->|AI Execution| AIAPI[Gemini External]\n    API <-->|Transactions| ORM[Django ORM]\n    ORM <--> SQL[(Database Store)]"
    add_mermaid_diagram(doc, comp_mermaid, "Figure 3.7: Component Diagram")

    add_heading(doc, '3.8 Deployment Diagram', level=2)
    add_bullet(doc, "The physical Deployment topography clarifies node configurations traversing hardware networking environments. A modern web browser securely interfaces across external TCP Ports tunneling inwards to fundamental Python application execution runtimes which internally manage localized disk reads representing the highly decoupled SQLite binary schema operations.")

    dep_mermaid = "flowchart TD\n    subgraph Client Device\n        Browser[Modern Internet Browser]\n    end\n    subgraph Server Hardware\n        WSGI[Gunicorn HTTP Server]\n        subgraph Application Container\n            APP[Django Backend]\n        end\n        subgraph Persistence Layer\n            SQL[(Relational Database)]\n        end\n    end\n    Browser -- HTTP Protocols --> WSGI\n    WSGI -- Sockets --> APP\n    APP -- I/O Operations --> SQL"
    add_mermaid_diagram(doc, dep_mermaid, "Figure 3.8: Deployment Diagram")

    add_heading(doc, '3.9 UML', level=2)
    add_bullet(doc, "The comprehensive Unified Modeling Language (UML) structural suite natively provided within this documentation guarantees an uncompromising, standardized, completely exhaustive, and extremely precise technical representation of Spendora.")
    add_bullet(doc, "It covers multi-dimensional architectural complexities simultaneously mapping chronological state progression, static storage architecture, operational client flows, class-based object representations, subsystem decoupling packages, and complete physical deployment configurations ensuring future engineering teams maintain a rigid universally understood standardized blueprint map.")

    add_heading(doc, '5. DRAWBACKS AND LIMITATIONS', level=1)
    add_bullet(doc, "The system intrinsically requires consistent functional internet backbone connectivity; any localized degradation or severe interruption strictly disables out-bound asynchronous requests halting Artificial Intelligence processing rendering receipt image extraction and audio interpretation modules entirely offline.")
    add_bullet(doc, "Given current structural choices leaning towards extreme localized decoupling, real-time sync connectivity with major institutional banking structures using Plaid or Yodlee vectors is prohibited, inherently meaning transaction feeds continue independently isolated relying purely on accurate manual engagement over extended lifestyle horizons.")
    add_bullet(doc, "Deep artificial intelligence visual interpretations exhibit extreme accuracy but remain subjectively prone to probabilistic hallucinations in specific circumstances where graphical resolution significantly lacks fidelity, meaning poorly lit receipt photographic input may occasionally map values improperly without rigid human overview validation prior towards finalized insertion.")

    add_heading(doc, '6. CONCLUSION', level=1)
    add_bullet(doc, "The execution and successful deployment of the Spendora Advanced Financial AI System represent an unparalleled architectural achievement merging simplistic human-center interface logic with immensely advanced deep learning intelligence parameters into an inherently comprehensive, secure tracker environment.")
    add_bullet(doc, "We have fundamentally achieved our core developmental objectives yielding an interactive capability allowing users intuitive and zero friction inputs resolving naturally into visually structured high-availability mathematical charts actively improving standard consumer lifestyle patterns natively.")

    add_heading(doc, '7. BIBLIOGRAPHY', level=1)
    add_bullet(doc, "1. Django Software Foundation, (2024). Django 4.x Official Documentation Core Framework Guidelines. Retrieved via djangoproject.com.")
    add_bullet(doc, "2. Google DeepMind AI Team, (2025). Google AI Studio API Endpoints and Prompt Design Implementation Protocol. Retrieved via ai.google.dev.")
    add_bullet(doc, "3. Mozilla Developer Network, HTML5 & Vanilla JavaScript Syntactical Specifications. Retrieved via developer.mozilla.org.")

    doc.save('Spendora_Project_Documentation.docx')

if __name__ == "__main__":
    main()
